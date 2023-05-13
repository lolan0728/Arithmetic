# & *********************************************
# & @Date: 2023-04-28 21:08:27
# & @LastEditors: lolan0728 vampire.lolan@outlook.com
# & @LastEditTime: 2023-05-12 22:02:31
# & @FilePath: /Arithmetic/IO/OutputToDocx.py
# & @Description: Wordファイルに出力
# & *********************************************
import os
import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


class OutputToDocx:

    # *********************************************
    # * @description: 初期化
    # * @param {str} title: ファイルのタイトル
    # * @param {str} savePath: 保存パス
    # * @param {str} date: 日付
    # *********************************************
    def __init__(self, title: str, savePath: str, date: str = None) -> None:
        # タイトル
        self.title = title
        # 保存パス
        self.savePath = savePath
        # 日付、設定しない場合当日を利用
        self.date = datetime.date.today().strftime("%Y/%m/%d") \
            if date is None else date

    # *********************************************
    # * @description: Wordファイルに出力
    # * @param {str} filename: ファイル名
    # * @param {list[str]} records: 数式リスト
    # *********************************************
    def output(self, filename: str, pageSize: int,
               formulas: list[str]) -> None:
        document = Document()
        section = document.sections[0]

        # ブランク設定
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(0.7)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        # タイトル
        title_ = document.add_heading(level=0)
        title_run = title_.add_run(self.title)
        title_run.font.size = Pt(24)
        title_run.font.name = 'Times New Roman'
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        title_run.bold = True

        # 日付
        date_ = document.add_paragraph()
        date_.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_.add_run(self.date)

        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(14)
        date_run.italic = True

        # 数式部分、１行４列
        table = document.add_table(rows=0, cols=4)
        table.autofit = True
        lstFormula = self.__splitTo4Col(formulas, pageSize)
        for a, b, c, d in lstFormula[0]:
            row = table.add_row()
            row_cells = row.cells

            p = row_cells[0].paragraphs[0]
            run = p.add_run(a)
            run.font.size = Pt(12)

            p = row_cells[1].paragraphs[0]
            run = p.add_run(b)
            run.font.size = Pt(12)

            p = row_cells[2].paragraphs[0]
            run = p.add_run(c)
            run.font.size = Pt(12)

            p = row_cells[3].paragraphs[0]
            run = p.add_run(d)
            run.font.size = Pt(12)

        path = os.path.join(self.savePath, filename)
        document.save(path)

    # *********************************************
    # * @description: 数式を1行4列に分割
    # * @param {list[str])} formulas: 数式文字列リスト
    # * @param {int} pageSize: 1ページの数式数量
    # * @return {list[list[str]]}: 分割したリスト
    # *********************************************
    def __splitTo4Col(self,
                      formulas: list[str],
                      pageSize: int = 100) -> list[list[str]]:
        result: list[list[str]] = []
        while len(formulas) > 0:
            pageSize = pageSize if len(formulas) >= pageSize else len(formulas)
            ls = formulas[:pageSize]
            formulas = formulas[pageSize:]
            rd = []
            while len(ls) > 0:
                rs = 4 if len(ls) > 3 else len(ls)
                row = ls[:rs]
                if rs < 4:
                    for _ in range(4 - rs):
                        row.append('')
                rd.append(row)
                ls = ls[rs:]
            result.append(rd)
        return result


# テスト用
if __name__ == "__main__":
    records = [('88 + 88 + 88 = 88', '88 + 88 + 88 = 88', '88 + 88 + 88 = 88',
                '88 + 88 + 88 = 88')] * 25
    params = {
        'title': "Harvey's Math Exercises",
        'savePath': '//Users//lolan//Desktop//Arithmetic//',
        'date': None
    }
    ins = OutputToDocx(**params)
    ins.output('test.docx', 100, records)
