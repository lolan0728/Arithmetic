# *********************************************
# * @Date: 2023-04-28 21:08:27
# * @LastEditors: lolan0728 vampire.lolan@outlook.com
# * @LastEditTime: 2023-05-06 16:25:12
# * @FilePath: /Arithmetic/IO/OutPutter.py
# * @Description: 数式出力
# *********************************************
import os
import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


# *********************************************
# * @description: 数式をWordファイルに出力
# * @Date: 2023-05-05 10:10:57
# *********************************************
class OutputDocx:
    def __init__(self, params: dict = None) -> None:
        self.setParams(params)

    # *********************************************
    # * @description: パラメーター設定
    # * @param {dict} params: パラメーターDictionary
    # * @return {*}: None
    # * @Date: 2023-05-05 10:16:20
    # *********************************************
    def setParams(self, params: dict) -> None:
        if params != None:
            # タイトル
            self.title = params['title']
            # 保存パス
            self.savePath = params['savePath']
            # 日付、設定しない場合当日を利用
            self.date = datetime.date.today().strftime("%Y/%m/%d") \
                if params['date'] is None else params['date']

    # *********************************************
    # * @description: Wordファイル出力
    # * @param {str} filename: ファイル名
    # * @param {list} records: 数式リスト
    # * @return {*}: None
    # * @Date: 2023-05-05 10:26:02
    # *********************************************
    def output(self, filename: str, records: list) -> None:
        document = Document()
        section = document.sections[0]

        # ブランク設定
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(0.7)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        # タイトル
        title_ = document.add_heading(level=0)
        title_run = title_.add_run(self.title)
        title_run.font.size = Pt(24)
        title_run.font.name = 'Times New Roman'
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        title_run.bold = True

        # 日付
        date_ = document.add_paragraph()
        date_.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_.add_run(self.date)

        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(14)
        date_run.italic = True

        # 数式部分、１行４列
        table = document.add_table(rows=0, cols=4)
        table.autofit = True
        for a, b, c, d in records:
            row = table.add_row()
            row_cells = row.cells

            p = row_cells[0].paragraphs[0]
            run = p.add_run(a)
            run.font.size = Pt(12)

            p = row_cells[1].paragraphs[0]
            run = p.add_run(b)
            run.font.size = Pt(12)

            p = row_cells[2].paragraphs[0]
            run = p.add_run(c)
            run.font.size = Pt(12)

            p = row_cells[3].paragraphs[0]
            run = p.add_run(d)
            run.font.size = Pt(12)

        path = os.path.join(self.savePath, filename)
        document.save(path)


# テスト用
if __name__ == "__main__":
    records = [('88 + 88 + 88 = 88', '88 + 88 + 88 = 88', '88 + 88 + 88 = 88',
                '88 + 88 + 88 = 88')] * 25
    params = {
        'title': "Harvey's Math Exercises",
        'savePath': '//Users//lolan//Desktop//Arithmetic//',
        'date': None
    }
    ins = OutputDocx(params)
    ins.output('test.docx', records)
