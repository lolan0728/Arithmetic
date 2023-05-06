
import os
import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt


class OutputDocx:
    def __init__(self, title, name, savepath, startdate=None):
        self.title = title
        self.name = name
        self.savepath = savepath
        self.startdate = datetime.date.today().strftime("%Y/%m/%d") \
             if startdate is None else startdate

    def output(self, filename, records, date=None):
        if date is None:
            date = self.startdate
            # date = ""

        document = Document()
        section = document.sections[0]
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

        title_ = document.add_heading(level=0)
        # title_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_.add_run(self.title)
        title_run.font.size = Pt(24)
        title_run.font.name = 'Times New Roman'
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        title_run.bold = True

        date_ = document.add_paragraph()
        date_.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_.add_run(date)

        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(14)
        date_run.italic = True

        table = document.add_table(rows=0, cols=4)
        table.autofit = True
        # table.style.font.size=Pt(14)
        for a, b, c, d in records:
            row = table.add_row()
            row_cells = row.cells

            p = row_cells[0].paragraphs[0]
            run = p.add_run(a)
            run.font.size = Pt(13)

            p = row_cells[1].paragraphs[0]
            run = p.add_run(b)
            run.font.size = Pt(13)

            p = row_cells[2].paragraphs[0]
            run = p.add_run(c)
            run.font.size = Pt(13)

            p = row_cells[3].paragraphs[0]
            run = p.add_run(d)
            run.font.size = Pt(13)
            # row_cells[0].text = a
            # row_cells[1].text = b
            # row_cells[2].text = c
            # row_cells[3].text = d

        path = os.path.join(self.savepath, filename)
        document.save(path)


if __name__ == "__main__":
    #self tester
    records = (
        ('17 + 18 + 1 = 36', '17 + 18 + 11 = 46', '17 + 18 + 1 = 36',
         '17 + 18 + 1 = 36'),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 = 36'),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 = 36'),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 = 36'),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
        ('17 + 18 + 1 =   ', '17 + 18 + 11 =   ', '17 + 18 + 1 =   ',
         '17 + 18 + 1 =   '),
    )
    OutputDocx('小八的算术练习', '周梓扬',
               'C:\\Users\\eos\\Desktop\\').output('test.docx', records)
